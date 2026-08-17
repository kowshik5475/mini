"""
Phase 3: DOCX template structure analysis.

Reads a .docx file with python-docx and extracts structure only --
title, headings (flat + nested hierarchy), paragraphs, tables, images,
basic formatting, and summary statistics. No AI, no content generation.
"""

from __future__ import annotations

import os

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

# Word style names we recognize as headings, mapped to a numeric level.
# Title/Subtitle are handled separately for document-title detection and
# are not part of the heading hierarchy.
HEADING_LEVELS = {
    "Heading 1": 1,
    "Heading 2": 2,
    "Heading 3": 3,
    "Heading 4": 4,
    "Heading 5": 5,
    "Heading 6": 6,
}

ALIGNMENT_NAMES = {
    0: "left",
    1: "center",
    2: "right",
    3: "justify",
}

# Cap how much paragraph text we send to the frontend for a single preview
# item, so one enormous paragraph can't bloat the response.
MAX_PARAGRAPH_PREVIEW_CHARS = 600


class DocxAnalysisError(Exception):
    """Raised when a DOCX file cannot be meaningfully analyzed."""

    def __init__(self, message: str, status_code: int = 422):
        super().__init__(message)
        self.message = message
        self.status_code = status_code


def _paragraph_alignment(paragraph) -> str | None:
    align = paragraph.alignment
    if align is None:
        return None
    return ALIGNMENT_NAMES.get(int(align), None)


def _paragraph_formatting(paragraph) -> dict:
    """Best-effort formatting from the paragraph's runs.

    A paragraph can contain multiple runs with different formatting;
    for Phase 3 we report the first run with visible text as a
    representative sample, and fall back to None for anything that
    isn't explicitly set (Word inherits a lot from styles/themes, and
    we don't attempt to resolve inheritance here).
    """
    bold = None
    italic = None
    font_size = None

    for run in paragraph.runs:
        if run.text.strip():
            bold = run.bold
            italic = run.italic
            font_size = run.font.size.pt if run.font.size else None
            break

    return {
        "bold": bold,
        "italic": italic,
        "font_size": font_size,
        "alignment": _paragraph_alignment(paragraph),
    }


def _extract_title(paragraphs) -> str | None:
    """Priority: a paragraph styled 'Title', else the first non-empty
    paragraph in the document. Never invents a title."""
    for p in paragraphs:
        if p.style and p.style.name == "Title" and p.text.strip():
            return p.text.strip()

    for p in paragraphs:
        if p.text.strip():
            return p.text.strip()

    return None


def _extract_headings(paragraphs) -> list[dict]:
    headings = []
    for p in paragraphs:
        style_name = p.style.name if p.style else None
        level = HEADING_LEVELS.get(style_name)
        if level is None:
            continue
        text = p.text.strip()
        if not text:
            continue
        formatting = _paragraph_formatting(p)
        headings.append(
            {
                "text": text,
                "level": level,
                "style": style_name,
                **formatting,
            }
        )
    return headings


def _build_hierarchy(headings: list[dict]) -> list[dict]:
    """Turn a flat, ordered heading list into a nested tree based on
    Word heading levels (not text like "1.1", "2.3" in the paragraph)."""
    root: list[dict] = []
    # stack entries: (level, children_list_to_append_into)
    stack: list[tuple[int, list[dict]]] = [(0, root)]

    for h in headings:
        node = {"title": h["text"], "level": h["level"], "children": []}
        while stack and stack[-1][0] >= h["level"]:
            stack.pop()
        stack[-1][1].append(node)
        stack.append((h["level"], node["children"]))

    return root


def _extract_paragraphs(paragraphs) -> list[dict]:
    result = []
    for p in paragraphs:
        text = p.text.strip()
        if not text:
            continue
        style_name = p.style.name if p.style else "Normal"
        truncated = text[:MAX_PARAGRAPH_PREVIEW_CHARS]
        result.append({"text": truncated, "style": style_name})
    return result


def _extract_tables(document) -> list[dict]:
    tables = []
    for index, table in enumerate(document.tables):
        rows = len(table.rows)
        columns = len(table.columns) if table.columns else (len(table.rows[0].cells) if rows else 0)

        headers = []
        if rows > 0:
            headers = [cell.text.strip() for cell in table.rows[0].cells]

        tables.append(
            {
                "index": index,
                "rows": rows,
                "columns": columns,
                "headers": headers,
            }
        )
    return tables


def _extract_images(document) -> list[dict]:
    images = []
    index = 0
    for rel in document.part.rels.values():
        if "image" not in rel.reltype:
            continue
        try:
            partname = str(rel.target_part.partname)
            content_type = rel.target_part.content_type
        except Exception:
            # Some rels (e.g. external/broken references) may not resolve
            # to a usable part -- skip rather than fail the whole analysis.
            continue

        filename = os.path.basename(partname)
        image_type = content_type.split("/")[-1] if content_type else _extension_of(filename)

        images.append({"index": index, "filename": filename, "type": image_type})
        index += 1

    return images


def _extension_of(filename: str) -> str:
    if "." not in filename:
        return "unknown"
    return filename.rsplit(".", 1)[-1].lower()


def analyze_docx(file_path: str, original_filename: str | None = None) -> dict:
    """Open and analyze a .docx file, returning a structured summary.

    Raises DocxAnalysisError for corrupted files or documents with no
    readable content, which the route layer turns into a clean HTTP error.
    """
    try:
        document = Document(file_path)
    except PackageNotFoundError as exc:
        raise DocxAnalysisError(
            "Unable to read this Word document. The file may be corrupted."
        ) from exc
    except Exception as exc:  # pragma: no cover - safety net for unexpected zip/xml errors
        raise DocxAnalysisError(
            "Unable to read this Word document. The file may be corrupted."
        ) from exc

    paragraphs = document.paragraphs

    title = _extract_title(paragraphs)
    headings = _extract_headings(paragraphs)
    structure = _build_hierarchy(headings)
    paragraph_previews = _extract_paragraphs(paragraphs)
    tables = _extract_tables(document)
    images = _extract_images(document)

    has_any_content = bool(paragraph_previews or headings or tables or images)
    if not has_any_content:
        raise DocxAnalysisError("No readable content was found in this document.")

    paragraph_count = len(paragraph_previews)
    heading_count = len(headings)
    table_count = len(tables)
    image_count = len(images)

    return {
        "filename": original_filename or os.path.basename(file_path),
        "title": title,
        "paragraph_count": paragraph_count,
        "heading_count": heading_count,
        "table_count": table_count,
        "image_count": image_count,
        "headings": headings,
        "structure": structure,
        "paragraphs": paragraph_previews,
        "tables": tables,
        "images": images,
        "statistics": {
            "paragraphs": paragraph_count,
            "headings": heading_count,
            "tables": table_count,
            "images": image_count,
            "sections": heading_count,
        },
    }